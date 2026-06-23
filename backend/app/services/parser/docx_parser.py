from pathlib import Path

from docx import Document

from app.schemas.parser import ParsedDocument
from app.services.parser.base import BaseParser


class DOCXParser(BaseParser):
    """
    Parser for DOCX files.
    """

    def extract_text(self, file_path: Path) -> ParsedDocument:

        document = Document(file_path)

        content = "\n".join(
            paragraph.text
            for paragraph in document.paragraphs
        )

        return ParsedDocument(
            filename=file_path.name,
            file_type="docx",
            content=content,
        )